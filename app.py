import os
import re
from io import BytesIO
from flask import Flask, request, send_file, jsonify, send_from_directory
from flask_cors import CORS
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet

# ------------------------------------------
# APP SETUP
# ------------------------------------------
app = Flask(__name__, static_folder=".")
CORS(app)

OPENAI_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_KEY:
    raise EnvironmentError("❌ Missing OPENAI_API_KEY — add it in Railway Variables")

client = OpenAI(api_key=OPENAI_KEY)

# ------------------------------------------
# HELPERS
# ------------------------------------------
def clean_ai_output(text: str) -> str:
    """Clean unwanted artifacts like Tailored CV or markdown lines."""
    text = re.sub(r"(?i)tailored\s*(cv|resume)", "", text)
    text = re.sub(r"[-_=]{2,}", "", text)
    return text.strip()

def create_docx(cv_text: str, target_name: str):
    """Generate a DOCX CV."""
    doc = Document()
    clean_title = target_name.strip().title().replace("_", " ")

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run(f"{clean_title} CV")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph()

    # Split by sections
    sections = re.split(
        r"(?=Summary|Key Skills|Professional Experience|Education|Certifications|Additional Information)",
        cv_text,
    )

    for section in sections:
        lines = section.strip().splitlines()
        if not lines:
            continue
        heading = lines[0].strip()
        content = "\n".join(lines[1:]).strip()

        doc.add_heading(heading, level=2)
        for para in content.split("\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(cv_text: str, target_name: str):
    """Generate a PDF CV."""
    buffer = BytesIO()
    clean_title = target_name.strip().title().replace("_", " ")
    styles = getSampleStyleSheet()

    pdf = SimpleDocTemplate(buffer, pagesize=A4)
    story = []

    story.append(Paragraph(f"<b>{clean_title} CV</b>", styles["Title"]))
    story.append(Spacer(1, 12))

    for line in cv_text.split("\n"):
        if re.match(r"^(Summary|Key Skills|Professional Experience|Education|Certifications|Additional Information)", line.strip()):
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"<b>{line.strip()}</b>", styles["Heading2"]))
        elif line.strip():
            story.append(Paragraph(line.strip(), styles["Normal"]))
        else:
            story.append(Spacer(1, 6))

    pdf.build(story)
    buffer.seek(0)
    return buffer

# ------------------------------------------
# ROUTES
# ------------------------------------------
@app.route("/")
def home():
    return send_from_directory(".", "index.html")

@app.route("/generate", methods=["POST"])
def generate_cv():
    try:
        data = request.get_json(force=True)
        old_cv = data.get("old_cv", "").strip()
        job_desc = data.get("job_desc", "").strip()
        target_name = (data.get("target_name") or "").strip()
        file_format = (data.get("file_format") or "docx").lower()

        if not old_cv or not job_desc or not target_name:
            return jsonify({"error": "Missing one or more fields: old_cv, job_desc, or target_name"}), 400

        print(f"🧩 Received target_name: {repr(target_name)} | file_format: {file_format}")

        # ----------------- FULL PROFESSIONAL PROMPT -----------------
        prompt = f"""
You are a highly skilled resume writer specializing in Applicant Tracking System (ATS) optimization and crafting compelling narratives for mid-career to executive-level professionals. 
Your task is to rewrite the provided CV to not only align with the given job description but also to significantly improve its impact and persuasiveness for a human reader. 
Prioritize clarity, quantifiable achievements expressed with powerful action verbs, and the strategic integration of relevant keywords to ensure the CV is both ATS-friendly and compelling.

The output should strictly adhere to the following structure, using either Tahoma, Arial, or Times New Roman font (though ultimately, plain text will be output): 
============================
FULL NAME
LOCATION | CONTACT INFO | EMAIL | LINKEDIN

Summary
[Concise 2–3 line summary designed to hook the reader, highlighting key strengths most relevant to the job.]

Key Skills
- [List 8–12 carefully chosen key skills with relevant keywords and measurable context.]

Professional Experience
[Company Name], [Job Title] | [Dates of Employment]
- [3–5 bullet points quantifying results and achievements using the STAR method.]

Education
[Degree], [Institution], [Graduation Year (or GPA if above 3.5)]

Certifications
[List relevant certifications]

Additional Information (Optional)
[Languages, Awards, Publications, or Tools directly relevant to the job.]
============================

INPUT:
JOB DESCRIPTION: {job_desc}
CURRENT CV: {old_cv}

OUTPUT:
Provide only the rewritten CV in plain text — no markdown, explanations, or filler text. 
Focus on precision, clarity, and measurable results.
"""

        # ----------------- OPENAI CALL -----------------
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )

        updated_cv = clean_ai_output(response.choices[0].message.content.strip())

        # ----------------- GENERATE FILE -----------------
        clean_name = target_name.strip().title()
        filename = f"{clean_name.replace(' ', '_')}.{file_format}"

        if file_format == "pdf":
            buffer = create_pdf(updated_cv, clean_name)
            mimetype = "application/pdf"
        else:
            buffer = create_docx(updated_cv, clean_name)
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        print(f"✅ Final filename: {filename}")

        return send_file(buffer, as_attachment=True, download_name=filename, mimetype=mimetype)

    except Exception as e:
        print("❌ Backend error:", e)
        return jsonify({"error": str(e)}), 500

@app.route("/health")
def health():
    return jsonify({"status": "ok", "message": "Backend running fine"})

# ------------------------------------------
# MAIN
# ------------------------------------------
if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
